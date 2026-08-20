from docx import Document

doc_path = r"c:\Users\DEMİRCİ\OneDrive\Masaüstü\entegrasyon\n11APISoapREFERANSDOKUMANTASYONU_v10_0.docx"
try:
    doc = Document(doc_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Let's search for "auth", "appKey", "appSecret" or "GetProductList"
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if "appKey" in text or "appSecret" in text or "<auth>" in text or "SOAP" in text:
            print(f"[Paragraph {i}] {text[:200]}")
            
except Exception as e:
    print(f"Error reading docx: {e}")
